import os
import tkinter as tk
from tkinter import *
from tkinter import Toplevel
from PIL import Image,ImageTk
from tkinter import messagebox, Button, Label, Frame, Entry, scrolledtext
import google.generativeai as genai
import speech_recognition as sr
import requests
from bs4 import BeautifulSoup
from docx import Document
from docx2pdf import convert



API_KEY = "" #buraya kendi api keyinizi yazarak programı çalıştırabilirsiniz
genai.configure(api_key=API_KEY)



def ögrenci_login():
    global ögrenci_kodu_entry
    global ögrenci_adsoyad_entry
    ögrenci_login = Toplevel(login)
    ögrenci_login.title("ÖĞRENCİ LOGIN")
    ögrenci_login.geometry("1080x570+380+210")
    ögrenci_login.configure(bg="light gray")

    #mavi kutu ve içi
    bluebox = tk.Label(ögrenci_login,bg="dark blue" ,width=80 ,height=570)
    bluebox.place(x=0,y=0)

    bluebox_text = tk.Label(ögrenci_login,text="EĞİTİM ÖĞRETİM PANELİ (EÖP)" ,font="Inter 20 bold" ,bg="dark blue" ,fg="white")
    bluebox_text.place(x=60,y=90)

    bluebox_text2 = tk.Label(ögrenci_login,text="EÖP,Öğretmenler ve Öğrenciler için yazılan \nve eğitim-öğretim için gerekli olabilecek bazı \nseçenekler sunan programdır" ,font="Inter 10 bold" ,bg="dark blue" ,fg="white")
    bluebox_text2.place(x=115,y=200)


    #mavi kutu yanı
    ögrenci_girisi_text = tk.Label(ögrenci_login,text="ÖĞRENCİ GİRİŞ" ,font="Inter 20 bold" ,bg="light gray" ,fg="dark blue")
    ögrenci_girisi_text.place(x=615,y=70)

    ögrenci_adsoyad = tk.Label(ögrenci_login,text="İsim ve Soyisim:" ,font="Inter 15 bold" ,bg="light gray" ,fg="dark blue")
    ögrenci_adsoyad.place(x=630,y=150)

    ögrenci_adsoyad_entry = tk.Entry(ögrenci_login ,font="Inter 15 bold" ,bg="white" ,width=30)
    ögrenci_adsoyad_entry.place(x=630,y=180)

    ögrenci_kodu = tk.Label(ögrenci_login,text="Öğrenci Kodu:" ,font="Inter 15 bold" ,bg="light gray" ,fg="dark blue")
    ögrenci_kodu.place(x=630,y=230)

    ögrenci_kodu_entry = tk.Entry(ögrenci_login ,font="Inter 15 bold" ,bg="white" ,width=30 ,show="*")
    ögrenci_kodu_entry.place(x=630,y=260)

    ögrenci_girisyapma_buton = tk.Button(ögrenci_login,text='Giriş Yap' ,font='Inter 16 bold' ,bg='blue' ,fg="white" ,command=ögrenci_check_password).place(x=850 ,y=315)

    ögretmen_girisegit_buton = tk.Button(ögrenci_login,text='Öğretmen Giriş' ,font='Inter 10 bold' ,bg='blue' ,fg="white" ,command=ögretmen_login).place(x=630 ,y=320)

    login.withdraw()






def ögretmen_login():
    global ögretmen_kodu_entry 
    global ögretmen_adsoyad_entry
    ögretmen_login = Toplevel(login)
    ögretmen_login.title("ÖĞRETMEN LOGIN")
    ögretmen_login.geometry("1080x570+380+210")
    ögretmen_login.configure(bg="light gray")

    #mavi kutu ve içi
    bluebox = tk.Label(ögretmen_login,bg="dark blue" ,width=80 ,height=570)
    bluebox.place(x=0,y=0)

    bluebox_text = tk.Label(ögretmen_login,text="EĞİTİM ÖĞRETİM PANELİ (EÖP)" ,font="Inter 20 bold" ,bg="dark blue" ,fg="white")
    bluebox_text.place(x=60,y=90)

    bluebox_text2 = tk.Label(ögretmen_login,text="EÖP,Öğretmenler ve Öğrenciler için yazılan \nve eğitim-öğretim için gerekli olabilecek bazı \nseçenekler sunan programdır" ,font="Inter 10 bold" ,bg="dark blue" ,fg="white")
    bluebox_text2.place(x=115,y=200)


    #mavi kutu yanı
    ögretmen_girisi_text = tk.Label(ögretmen_login,text="ÖĞRETMEN GİRİŞ" ,font="Inter 20 bold" ,bg="light gray" ,fg="dark blue")
    ögretmen_girisi_text.place(x=615,y=70)

    ögretmen_adsoyad = tk.Label(ögretmen_login,text="İsim ve Soyisim:" ,font="Inter 15 bold" ,bg="light gray" ,fg="dark blue")
    ögretmen_adsoyad.place(x=630,y=150)

    ögretmen_adsoyad_entry = tk.Entry(ögretmen_login ,font="Inter 15 bold" ,bg="white" ,width=30)
    ögretmen_adsoyad_entry.place(x=630,y=180)

    ögretmen_kodu = tk.Label(ögretmen_login,text="Öğretmen Kodu:" ,font="Inter 15 bold" ,bg="light gray" ,fg="dark blue")
    ögretmen_kodu.place(x=630,y=230)

    ögretmen_kodu_entry = tk.Entry(ögretmen_login ,font="Inter 15 bold" ,bg="white" ,width=30 ,show="*")
    ögretmen_kodu_entry.place(x=630,y=260)

    ögretmen_girisyapma_buton = tk.Button(ögretmen_login,text='Giriş Yap' ,font='Inter 16 bold' ,bg='blue' ,fg="white" ,command=ögretmen_check_password).place(x=850 ,y=315)

    ögrenci_girisegit_buton = tk.Button(ögretmen_login,text='Öğrenci Giriş' ,font='Inter 10 bold' ,bg='blue' ,fg="white" ,command=ögrenci_login).place(x=630 ,y=320)

    login.withdraw()



def ögretmen_check_password():
    password = ögretmen_kodu_entry.get()
    adsoyad = ögretmen_adsoyad_entry.get()
    if password == "12345":
        ögretmen_panel()
        messagebox.showinfo("Giriş Başarılı", f"Hoş geldin, {adsoyad}!")
    else:
        messagebox.showerror("Hata", "Şifre yanlış!")





def ögrenci_check_password():
    password = ögrenci_kodu_entry.get()
    adsoyad = ögrenci_adsoyad_entry.get()
    if password == "123":  
        ögrenci_panel()
        messagebox.showinfo("Giriş Başarılı", f"Hoş geldin, {adsoyad}!")
    else:
        messagebox.showerror("Hata", "Şifre yanlış!")







def ögretmen_panel():
    global robot_tk_resmi
    global mikrofon_tk_resmi
    global sorucevap_tk_resmi
    global tarihler_tk_resmi

    global gif_tk_resmi

    ögretmen_panel = Toplevel(login)
    ögretmen_panel.title("ÖĞRETMEN PANEL")
    ögretmen_panel.geometry("1080x570+380+210")
    ögretmen_panel.configure(bg="white")




    #arka plan box resim
    gif_resmi = Image.open("ögretmenboximage.gif") 
    gif_resmi = gif_resmi.resize((1080, 570), Image.LANCZOS) 
    gif_tk_resmi = ImageTk.PhotoImage(gif_resmi)




    gifbox = tk.Label(ögretmen_panel,image=gif_tk_resmi ,width=1080)
    gifbox.place(x=0,y=0)


    ögretmen_page = tk.Label(ögretmen_panel,text="ÖĞRETMEN PAGE (EÖP)" ,font="Inter 30 bold" ,bg="#4696FF" ,fg="dark blue")
    ögretmen_page.place(x=300,y=50)



    #Resimlibuton (Resimler) ögretmen
    robot_resmi = Image.open("robot character.jpeg")  
    robot_resmi = robot_resmi.resize((50, 50), Image.LANCZOS)  
    robot_tk_resmi = ImageTk.PhotoImage(robot_resmi)

    mikrofon_resmi = Image.open("mikrofon.jpeg")  
    mikrofon_resmi = mikrofon_resmi.resize((50, 50), Image.LANCZOS) 
    mikrofon_tk_resmi = ImageTk.PhotoImage(mikrofon_resmi)

    sorucevap_resmi = Image.open("sorucevap.jpeg")  
    sorucevap_resmi = sorucevap_resmi.resize((50, 45), Image.LANCZOS) 
    sorucevap_tk_resmi = ImageTk.PhotoImage(sorucevap_resmi)

    tarihler_resmi = Image.open("tarihler.jpeg")  
    tarihler_resmi = tarihler_resmi.resize((50, 45), Image.LANCZOS)  
    tarihler_tk_resmi = ImageTk.PhotoImage(tarihler_resmi),


    #öğretmen panel butonlar
    eöpbot_buton = tk.Button(ögretmen_panel,text='EÖP BOT',image=robot_tk_resmi, compound=tk.LEFT,font='Inter 16 bold' ,bg='blue' ,fg="white" ,width=145 ,command=open_eöp_bot).place(x=100 ,y=250)
    seslikonus_buton = tk.Button(ögretmen_panel,text='Sesli Metin' ,image=mikrofon_tk_resmi, compound=tk.LEFT,font='Inter 16 bold' ,bg='blue' ,fg="white" ,width=160,command=open_voice_input).place(x=480 ,y=250)
    tarihpaylas_buton = tk.Button(ögretmen_panel,text='Tarih Paylaş' ,image=tarihler_tk_resmi, compound=tk.LEFT,font='Inter 16 bold' ,bg='blue' ,fg="white" ,width=175,command=open_write_dates).place(x=820 ,y=250)
    soruyaz_buton = tk.Button(ögretmen_panel,text='Soru Paylaş' ,image=sorucevap_tk_resmi, compound=tk.LEFT,font='Inter 16 bold' ,bg='blue' ,fg="white" ,width=170,command=open_write_question).place(x=100 ,y=400)





def ögrenci_panel():
    global robot_tk_resmi
    global sorucevap_tk_resmi
    global tarihler_tk_resmi
    global bilgial_tk_resmi

    global gif_tk_resmi


    ögrenci_panel = Toplevel(login)
    ögrenci_panel.title("ÖĞRENCİ PANEL")
    ögrenci_panel.geometry("1080x570+380+210")
    ögrenci_panel.configure(bg="white")


    #arka plan box resim
    gif_resmi = Image.open("gifboxögrenci.jpeg") 
    gif_resmi = gif_resmi.resize((1080, 570), Image.LANCZOS) 
    gif_tk_resmi = ImageTk.PhotoImage(gif_resmi)




    gifbox = tk.Label(ögrenci_panel,image=gif_tk_resmi ,width=1080)
    gifbox.place(x=0,y=0)


    ögrenci_page = tk.Label(ögrenci_panel,text="ÖĞRENCİ PAGE (EÖP)" ,font="Inter 30 bold" ,bg="light blue" ,fg="dark blue")
    ögrenci_page.place(x=320,y=50)





    #Resimlibuton (Resimler) ögrenci
    robot_resmi = Image.open("robot character.jpeg")  
    robot_resmi = robot_resmi.resize((50, 50), Image.LANCZOS)  
    robot_tk_resmi = ImageTk.PhotoImage(robot_resmi)

    sorucevap_resmi = Image.open("sorucevap.jpeg")  
    sorucevap_resmi = sorucevap_resmi.resize((50, 45), Image.LANCZOS)  
    sorucevap_tk_resmi = ImageTk.PhotoImage(sorucevap_resmi)

    tarihler_resmi = Image.open("tarihler.jpeg")  
    tarihler_resmi = tarihler_resmi.resize((50, 45), Image.LANCZOS)  
    tarihler_tk_resmi = ImageTk.PhotoImage(tarihler_resmi)

    bilgial_resmi = Image.open("bilgial.jpeg") 
    bilgial_resmi = bilgial_resmi.resize((50, 45), Image.LANCZOS)  
    bilgial_tk_resmi = ImageTk.PhotoImage(bilgial_resmi)


    #öğrenci panel butonlar
    eöpbot_buton = tk.Button(ögrenci_panel,text='EÖP BOT' ,image=robot_tk_resmi, compound=tk.LEFT,font='Inter 16 bold' ,bg='blue' ,fg="white" ,width=145 ,command=open_eöp_bot).place(x=100 ,y=250)
    önemlitarihler_buton = tk.Button(ögrenci_panel,text='Önemli Tarihler' ,image=tarihler_tk_resmi, compound=tk.LEFT,font='Inter 16 bold' ,bg='blue' ,fg="white" ,width=205,command=open_pdff).place(x=420 ,y=250)
    bilgilen_buton = tk.Button(ögrenci_panel,text='Günlük Bilgi' ,image=bilgial_tk_resmi, compound=tk.LEFT,font='Inter 16 bold' ,bg='blue' ,fg="white" ,width=173,command=web_data).place(x=780 ,y=250)
    sorucöz_buton = tk.Button(ögrenci_panel,text='Soru Çöz' ,image=sorucevap_tk_resmi, compound=tk.LEFT,font='Inter 16 bold' ,bg='blue' ,fg="white" ,width=143,command=open_pdf).place(x=100 ,y=400)






def open_eöp_bot():
    eöp_bot_window = tk.Toplevel(login)
    eöp_bot_window.geometry("1500x700+210+100")
    eöp_bot_window.title("EÖP BOT")
    eöp_bot_window.configure(bg="#10212D")

    # Başlık
    title_label = Label(eöp_bot_window, text="EÖP BOT", fg="white", bg="#10212D", font=("Arial", 24, "bold"))
    title_label.pack(pady=20)

    # Chat alanı
    chat_display = scrolledtext.ScrolledText(eöp_bot_window, wrap=tk.WORD, state='disabled', bg="#222222", fg="white")
    chat_display.pack(pady=10, padx=10, fill=tk.BOTH, expand=True)

    # Kullanıcı girişi
    user_entry = Entry(eöp_bot_window, width=100)
    user_entry.pack(pady=10)

    # Mesaj gönderme fonksiyonu
    def send_message():
        user_input = user_entry.get()
        if user_input:
            chat_display.configure(state='normal')
            chat_display.insert(tk.END, f"Sen: {user_input}\n")
            user_entry.delete(0, tk.END)

            # Model yapılandırması
            generation_config = {
                "temperature": 1,
                "top_p": 0.95,
                "top_k": 64,
                "max_output_tokens": 8192,
                "response_mime_type": "text/plain",
            }

            model = genai.GenerativeModel(
                model_name="gemini-1.5-pro",
                generation_config=generation_config,
            )

            # Chat oturumu başlat
            chat_session = model.start_chat(history=[])

            # Mesajı gönder
            response = chat_session.send_message(user_input)

            # Yanıtı göster
            chat_display.insert(tk.END, f"Bot: {response.text}\n")
            chat_display.configure(state='disabled')

    # Gönder butonu
    send_button = Button(eöp_bot_window, text="Gönder", command=send_message, bg="#5B5B5B", fg="white" ,font=("Arial", 10, "bold"))
    send_button.pack(pady=5)

    # Çıkış butonu
    exit_button = Button(eöp_bot_window, text="Kapat", command=eöp_bot_window.destroy, bg="#5B5B5B", fg="white", font=("Arial", 10, "bold"))
    exit_button.pack(pady=20)







def open_voice_input():
    voice_window = tk.Toplevel(login)
    voice_window.geometry("1500x700+210+100")
    voice_window.title("Sesli Metin Girişi")
    voice_window.configure(bg="#10212D")

    # Başlık
    title_label = Label(voice_window, text="Sesli Metin Girişi", fg="white", bg="#10212D", font=("Arial", 24, "bold"))
    title_label.pack(pady=20)

    # Scrollable metin alanı
    chat_display = scrolledtext.ScrolledText(voice_window, wrap=tk.WORD, state='disabled', bg="#222222", fg="white")
    chat_display.pack(pady=10, padx=10, fill=tk.BOTH, expand=True)

    # Durum etiketi
    label_status = Label(voice_window, text="", font=("Arial", 12), bg="#10212D", fg="white")
    label_status.pack(pady=5)

    # Sesli giriş fonksiyonu
    def start_listening():
        recognizer = sr.Recognizer()
        with sr.Microphone() as source:
            label_status.config(text="Dinliyor...")
            audio = recognizer.listen(source)
            label_status.config(text="Dinleme tamamlandı!")
            
            try:
                # Ses kaydını metne dönüştür
                text = recognizer.recognize_google(audio, language='tr-TR')
                chat_display.configure(state='normal')
                chat_display.insert(tk.END, f"Siz: {text}\n")
                chat_display.configure(state='disabled')
            except sr.UnknownValueError:
                chat_display.configure(state='normal')
                chat_display.insert(tk.END, "Bot: Ses anlaşılamadı.\n")
                chat_display.configure(state='disabled')
            except sr.RequestError as e:
                chat_display.configure(state='normal')
                chat_display.insert(tk.END, "Bot: API isteği başarısız oldu.\n")
                chat_display.configure(state='disabled')

    # Dinleme butonu
    listen_button = Button(voice_window, text="Mikrofonu Aç", command=start_listening, bg="#5A4B97", fg="white", font=("Arial", 12))
    listen_button.pack(pady=20)








#soruyaz yeri
def create_or_open_word():
    file_path = 'Soru_Sayfası.docx'
    
    if not os.path.exists(file_path):
        doc = Document()
        doc.add_heading('GÜNLÜK SORULAR', level=1)
        doc.add_paragraph('Her gün yeni sorularla eğitimde aktif kalın !')
        doc.save(file_path)
    
    return file_path

def add_content_to_word():
    file_path = create_or_open_word()
    doc = Document(file_path)

    # Text alanındaki içeriği al
    new_content = content_text.get("1.0", tk.END)  # 1. satırdan sonuna kadar al
    if new_content.strip():  # Boş değilse
        doc.add_paragraph(new_content.strip())
        doc.save(file_path)
        messagebox.showinfo("Bilgi", "Yeni içerik başarıyla eklendi!")
        
        # Belgeyi aç
        try:
            os.startfile(file_path)  # Windows için
        except Exception as e:
            messagebox.showerror("Hata", f"Belge açılamadı: {e}")
    else:
        messagebox.showwarning("Uyarı", "Lütfen içerik girin.")

def reset_page():
    file_path = 'Soru_Sayfası.docx'
    if os.path.exists(file_path):
        os.remove(file_path)
    create_or_open_word()
    content_text.delete("1.0", tk.END)  # Tüm metni temizle
    messagebox.showinfo("Bilgi", "Sayfa sıfırlandı.")

def convert_to_pdf():
    file_path = 'Soru_Sayfası.docx'
    pdf_path = 'Soru_Sayfası.pdf'
    convert(file_path, pdf_path)  # Word'ü PDF'ye dönüştür
    messagebox.showinfo("Bilgi", "Belge PDF formatına dönüştürüldü.")

def open_write_question():
    write_question = tk.Toplevel(login)  # login yerine root kullan
    write_question.configure(bg='light gray')
    write_question.title("Soru Yazma Alanı")
    write_question.geometry("1080x570+380+210")

    # Kaydırılabilir metin alanı oluştur
    global content_text
    content_text = Text(write_question, wrap=tk.WORD, width=100, height=20)
    content_text.pack(pady=10)

    # Scrollbar oluştur
    scrollbar = Scrollbar(write_question, command=content_text.yview)
    scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
    content_text.config(yscrollcommand=scrollbar.set)

    add_button = tk.Button(write_question, text="Word Belgesine Ekle", command=add_content_to_word)
    add_button.pack(pady=10)

    reset_button = tk.Button(write_question, text="Sayfayı Sıfırla", command=reset_page)
    reset_button.pack(pady=10)

    convert_button = tk.Button(write_question, text="Word'ü PDF'ye Dönüştür", command=convert_to_pdf)
    convert_button.pack(pady=10)


#pdf i aç
def open_pdf():
    pdf_path = 'Soru_Sayfası.pdf'
    try:
        os.startfile(pdf_path)  # Windows için
    except Exception as e:
        print(f"PDF açılamadı: {e}")









#Tarihyaz yeri
def create_or_open_wordd():
    file_path = 'Tarih_Sayfası.docx'
    
    if not os.path.exists(file_path):
        doc = Document()
        doc.add_heading('ÖNEMLİ TARİHLER', level=1)
        doc.add_paragraph('GÜNCEL SINAV,QUİZ vb Tarihlerine buradan ulaşabilirsiniz !')
        doc.save(file_path)
    
    return file_path

def add_content_to_wordd():
    file_path = create_or_open_wordd()
    doc = Document(file_path)

    # Text alanındaki içeriği al
    new_content = content_textt.get("1.0", tk.END)  # 1. satırdan sonuna kadar al
    if new_content.strip():  # Boş değilse
        doc.add_paragraph(new_content.strip())
        doc.save(file_path)
        messagebox.showinfo("Bilgi", "Yeni içerik başarıyla eklendi!")
        
        # Belgeyi aç
        try:
            os.startfile(file_path)  # Windows için
        except Exception as e:
            messagebox.showerror("Hata", f"Belge açılamadı: {e}")
    else:
        messagebox.showwarning("Uyarı", "Lütfen içerik girin.")

def reset_pagee():
    file_path = 'Tarih_Sayfası.docx'
    if os.path.exists(file_path):
        os.remove(file_path)
    create_or_open_wordd()
    content_textt.delete("1.0", tk.END)  # Tüm metni temizle
    messagebox.showinfo("Bilgi", "Sayfa sıfırlandı.")

def convert_to_pdff():
    file_path = 'Tarih_Sayfası.docx'
    pdf_path = 'Tarih_Sayfası.pdf'
    convert(file_path, pdf_path)  # Word'ü PDF'ye dönüştür
    messagebox.showinfo("Bilgi", "Belge PDF formatına dönüştürüldü.")

def open_write_dates():
    write_date = tk.Toplevel(login)  # login yerine root kullan
    write_date.configure(bg='gray')
    write_date.title("Tarih Yazma Alanı")
    write_date.geometry("1080x570+380+210")

    # Kaydırılabilir metin alanı oluştur
    global content_textt
    content_textt = Text(write_date, wrap=tk.WORD, width=100, height=20)
    content_textt.pack(pady=10)

    # Scrollbar oluştur
    scrollbar = Scrollbar(write_date, command=content_textt.yview)
    scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
    content_textt.config(yscrollcommand=scrollbar.set)

    add_button = tk.Button(write_date, text="Word Belgesine Ekle", command=add_content_to_wordd)
    add_button.pack(pady=10)

    reset_button = tk.Button(write_date, text="Sayfayı Sıfırla", command=reset_pagee)
    reset_button.pack(pady=10)

    convert_button = tk.Button(write_date, text="Word'ü PDF'ye Dönüştür", command=convert_to_pdff)
    convert_button.pack(pady=10)


#pdf i aç
def open_pdff():
    pdf_path = 'Tarih_Sayfası.pdf'
    try:
        os.startfile(pdf_path)  # Windows için
    except Exception as e:
        print(f"PDF açılamadı: {e}")







#öğrenci veri bilgi alma
def web_data():
    url = "https://shiftdelete.net/"  # Güncel teknolojik bilgi veren bir site (shiftdelete)
    try:
        response = requests.get(url)
        response.raise_for_status()  # HTTP hatalarını kontrol et
        soup = BeautifulSoup(response.text, 'html.parser')

        headlines = soup.find_all('h2')
        info_window = tk.Toplevel(login)
        info_window.geometry("1080x570+380+210")
        info_window.title("Güncel Bilgiler")
        info_window.configure(bg="#10212D")

        # Scrollable metin alanı
        info_display = scrolledtext.ScrolledText(info_window, wrap=tk.WORD, state='normal', bg="#222222", fg="white")
        info_display.pack(pady=10, padx=10, fill=tk.BOTH, expand=True)

        for headline in headlines:
            info_display.insert(tk.END, headline.get_text() + "\n")
        info_display.configure(state='disabled')

    except requests.RequestException as e:
        messagebox.showerror("Hata", f"Veri çekme sırasında hata oluştu: {e}")






#ana sayfa(giriş yapma kısmı)
login = tk.Tk()
login.geometry("1080x570+380+210")
login.title("LOGIN Page")
login.configure(bg="light gray")


#mavi kutu ve içi
bluebox = tk.Label(login,bg="dark blue" ,width=80 ,height=570)
bluebox.place(x=0,y=0)

bluebox_text = tk.Label(login,text="EĞİTİM ÖĞRETİM PANELİ (EÖP)" ,font="Inter 20 bold" ,bg="dark blue" ,fg="white")
bluebox_text.place(x=60,y=90)

bluebox_text2 = tk.Label(login,text="EÖP,Öğretmenler ve Öğrenciler için yazılan \nve eğitim-öğretim için gerekli olabilecek bazı \nseçenekler sunan programdır" ,font="Inter 10 bold" ,bg="dark blue" ,fg="white")
bluebox_text2.place(x=115,y=200)

bluebox_text3 = tk.Label(login,text="FTS DEVELOPMENT-FURKAN TAHA SARANDA-2024" ,font="Inter 10 bold" ,bg="dark blue" ,fg="white")
bluebox_text3.place(x=115,y=500)



#anasayfa butonlar
#butonlar
ögrenci_giris_buton = tk.Button(login,text='ÖĞRENCİ GİRİŞ' ,font='Inter 30 bold' ,bg='blue' ,fg="white" ,width=16,command=ögrenci_login).place(x=630 ,y=170)


ögretmen_giris_buton = tk.Button(login,text='ÖĞRETMEN GİRİŞ' ,font='Inter 30 bold' ,bg='blue' ,fg="white" ,width=16,command=ögretmen_login).place(x=630 ,y=270)



login.mainloop()
